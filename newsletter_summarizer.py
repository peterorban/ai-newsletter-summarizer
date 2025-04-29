import os
import pickle
import base64
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from dotenv import load_dotenv
import openai
from bs4 import BeautifulSoup
from docx import Document

# Load environment variables
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']

def get_gmail_service():
    creds = None
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)
    return build('gmail', 'v1', credentials=creds)

def get_newsletter_emails(service, query='label:AI_News'):
    results = service.users().messages().list(userId='me', q=query).execute()
    messages = results.get('messages', [])
    emails = []
    for message in messages:
        msg = service.users().messages().get(userId='me', id=message['id'], format='full').execute()
        emails.append(msg)
    return emails

def extract_text_from_email(email):
    payload = email['payload']
    parts = payload.get('parts', [payload])
    text = ''
    for part in parts:
        if part['mimeType'] == 'text/plain':
            data = part['body']['data']
            text += base64.urlsafe_b64decode(data).decode()
        elif part['mimeType'] == 'text/html':
            data = part['body']['data']
            html = base64.urlsafe_b64decode(data).decode()
            soup = BeautifulSoup(html, 'html.parser')
            text += soup.get_text()
    return text

from transformers import pipeline

from transformers import pipeline

def summarize_text(text, max_chunk_length=1000, final_summary_max_length=500):
    """Summarizes long text by chunking it and then summarizing the summaries."""
    summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
    text_chunks = [text[i:i + max_chunk_length] for i in range(0, len(text), max_chunk_length)]
    summaries = []
    for chunk in text_chunks:
        try:
            input_length = len(chunk.split())  # Estimate input length by word count
            max_len = int(input_length * 0.6)  # Set max_length to roughly 60% of input length
            max_len = min(max_len, 200)       # Ensure max_length doesn't exceed a reasonable limit
            min_len = int(max_len * 0.3)      # Set min_length to roughly 30% of max_length
            summary = summarizer(chunk, max_length=max_len, min_length=min_len)[0]['summary_text']
            summaries.append(summary)
        except Exception as e:
            print(f"Error summarizing chunk: {e}")
            summaries.append("")
    combined_summary = " ".join(summaries)
    # Summarize the combined summaries.
    final_summary = summarizer(combined_summary, max_length=final_summary_max_length, min_length=50)[0]['summary_text']
    return final_summary

def create_summary_document(summaries, filename='newsletter_summary.docx'):
    document = Document()
    document.add_heading('Newsletter Summaries', 0)
    for summary in summaries:
        document.add_paragraph(summary)
    document.save(filename)

def main():
    service = get_gmail_service()
    emails = get_newsletter_emails(service)
    summaries = []
    for email in emails:
        text = extract_text_from_email(email)
        summary = summarize_text(text)
        summaries.append(summary)
    create_summary_document(summaries)
    print("Summaries created in newsletter_summary.docx")

if __name__ == '__main__':
    main()
